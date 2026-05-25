from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return p

def add_paragraph_custom(doc, text, bold=False, italic=False, first_line_indent=Cm(1.25)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table_from_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = col_widths[i]
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            if col_widths and i < len(col_widths):
                row_cells[i].width = col_widths[i]
    
    doc.add_paragraph()
    return table

doc = Document()
section = doc.sections[0]
section.page_height = Cm(27.94)
section.page_width = Cm(21.59)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(2.54)

# ========================
# SPRINT 2
# ========================
add_heading_custom(doc, '2.6.1.2.\tSprint 2', level=2)

add_heading_custom(doc, '2.6.1.2.1.\tObjetivo del sprint', level=3)
add_paragraph_custom(doc, 'Dotar a la simulación del verbo activo central: la acción de conectar nodos. Implementar la lógica que forja aristas dinámicamente validadas por medio de la "Matriz de Adyacencia", mientras devela controles paramétricos de retroalimentación en la interfaz como el manejo prudente de presupuesto "Dinero" (basado en el peso de las aristas).')

add_heading_custom(doc, '2.6.1.2.2.\tPlanificación del sprint', level=3)
add_paragraph_custom(doc, 'En este Sprint se desarrolló la esencia algorítmica de la conexión (MatrizAdyacencia.server.lua y GrafoHelpers.lua). Se dispuso un mecanismo que supervisa toda tentativa de conexión, analizando su legalidad frente a preceptos como la de "grafo dirigido o no dirigido", al tiempo que realiza las sustracciones lógicas en UI por concepto del costo operacional ("pseudopesos" formados por el trazado de cableado). La tabla 2.15 presenta las historias de usuario realizadas en el Sprint 2.')

add_paragraph_custom(doc, 'Tabla 2.15 Sprint Backlog - Sprint 2', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Descripción', 'Estimación (Horas)'],
    [
        ['TG 01', 'Crear conexiones', 'Permitir la selección secuencial clic-sobre-clic entre un nodo de origen y destino para crear una "Arista Física (Cable)".', '12'],
        ['TG 03', 'Eliminar conexiones', 'Crear una función de desmontaje controlado que barra tramos incorrectos suprimiéndolos del registro transaccional lógico de memoria.', '6'],
        ['TG 05', 'Validación de reglas', 'Asegurar la validación topológica negando flujos ilegales (reversas en caminos dirigidos) o superposición reiterativa de la misma arista.', '10'],
        ['TG 04', 'Gestión de pesos y presupuesto', 'Manejo contable del peso de las aristas, afectando de forma restada la variable económica base o presupuesto.', '10'],
        ['CL 007', 'HUD de misiones', 'Instalación de las cabeceras UI o paneles (PanelMisionesHUD.lua) que notifiquen en vivo las modificaciones del presupuesto ante cada operación.', '14'],
        ['Total', '', '', '52']
    ]
)

add_heading_custom(doc, '2.6.1.2.3.\tRevisión del sprint', level=3)
add_paragraph_custom(doc, 'La Tabla 2.16 Revisión de criterios de aceptación - Sprint 2 presenta la retrospectiva del sprint 2, donde se presentan las historias de usuario y los criterios de aceptación verificados.')

add_paragraph_custom(doc, 'Tabla 2.16 Revisión de criterios de aceptación - Sprint 2', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Criterios de Aceptación', 'Estado'],
    [
        ['TG 01', 'Crear conexiones', 'Existe una representación en forma de cable al seleccionar dos puntos y se inscribe un nuevo valor relacional (X|Y) exitosamente en la Matriz de Adyacencia centralizada.', 'Cumplido'],
        ['TG 03', 'Eliminar conexiones', 'Retiro eficaz de la entidad visual 3D y eliminación de las relaciones recíprocas en las entidades del estado simulado.', 'Cumplido'],
        ['TG 05', 'Validación de reglas', 'La aplicación imposibilita empíricamente la producción de aristas que no obedezcan relaciones booleanas o dirían estar prohibidas (flechas univia).', 'Cumplido'],
        ['TG 04', 'Gestión de pesos y presupuesto', 'No puede erigirse ningún cable si esta acción deja a la cuenta restando de cero, asumiendo su peso respectivo.', 'Cumplido'],
        ['CL 007', 'HUD de misiones', 'Cualquier adicción o resta nodal impacta velozmente y persistentemente frente a sus ojos en la lista.', 'Cumplido']
    ]
)

add_heading_custom(doc, '2.6.1.2.4.\tResultado del Sprint', level=3)
add_paragraph_custom(doc, 'Se logró consolidar la experiencia táctica. El usuario puede ahora tomar la iniciativa computacional tendiendo cables, errando pragmáticamente, borrando y experimentando las restricciones que los costes traen al presupuesto límite. Cada acción propaga un reordenamiento constante de toda la interconexión mediante una matriz lógica invisible al ojo pero precisa en sus números. Ver')

# ========================
# SPRINT 3
# ========================
add_heading_custom(doc, '2.6.1.3.\tSprint 3', level=2)

add_heading_custom(doc, '2.6.1.3.1.\tObjetivo del sprint', level=3)
add_paragraph_custom(doc, 'Otorgar un paquete de validación algorítmica profunda al jugador (Minimapa, Pestaña Analítica Inspector) a través del cual logre solicitar al sistema una evaluación por BFS/DFS para medir posibles casos de nodos huérfanos o componentes de red sin alcanzar. Completar esto estableciendo localmente la persistencia segura de su trayectoria de victorias como "Data Store".')

add_heading_custom(doc, '2.6.1.3.2.\tPlanificación del sprint', level=3)
add_paragraph_custom(doc, 'Este Sprint incrustó una "lupa diagnóstica" matemática al Gameplay del usuario. Se habilitaron modos de visibilidad abstracta de las topologías (Modalidad Radar Dron (M) y Analizador de Problemas de Red (Tab)). Posteriormente se anexó a este motor la calculadora ServicioGrafosAnalisis con Algoritmos de Búsqueda capaces de dictaminar de manera certera el triunfo o desastre del estudiante. La tabla 2.17 presenta las historias de usuario realizadas en el Sprint 3.')

add_paragraph_custom(doc, 'Tabla 2.17 Sprint Backlog - Sprint 3', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Descripción', 'Estimación (Horas)'],
    [
        ['CL 004', 'Validación algorítmica (BFS)', 'Máquina Validadora (Analítica BFS): Computa por barrido si el ensamblaje de red conforma los subgrafos requeridos omitiendo caídas geográficas o nodos aislados no deseados.', '12'],
        ['TG 06', 'Minimapa e inspector', 'Modalidad visual de radar y cuadro matriz inspector, indicando desconexiones críticas mediante color rojizo y aciertos en cyan.', '16'],
        ['CL 011', 'Retroalimentación de errores', 'Mecanismos semánticos de prevención y error local que anuncien con precisión sintáctica en dónde y qué nodo falló dentro de la validación matemática.', '8'],
        ['CL 008', 'Persistencia de progreso', 'Módulo ServicioProgreso y guardado local mediante Roblox DataStore para atesorar de forma eterna internamente puntaje, tiempo record y nivel victorioso.', '14'],
        ['Total', '', '', '50']
    ]
)

add_heading_custom(doc, '2.6.1.3.3.\tRevisión del sprint', level=3)
add_paragraph_custom(doc, 'La Tabla 2.18 Revisión de criterios de aceptación - Sprint 3 presenta la retrospectiva del sprint 3, donde se presentan las historias de usuario y los criterios de aceptación verificados.')

add_paragraph_custom(doc, 'Tabla 2.18 Revisión de criterios de aceptación - Sprint 3', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Criterios de Aceptación', 'Estado'],
    [
        ['CL 004', 'Validación algorítmica (BFS)', 'El motor BFS devuelve inequívocamente la advertencia "Nodos Aislados" cuando se abortan puentes cardinales de la red objetivo.', 'Cumplido'],
        ['TG 06', 'Minimapa e inspector', 'El minimapa refleja virtual y bidimensionalmente la misma geometría relacional estática que la red 3D, sin desfasarse visualmente en sus nodos.', 'Cumplido'],
        ['CL 011', 'Retroalimentación de errores', 'El inspector o GUI de alertas traduce en términos del "sabor de mundo" (Jefe/Ciudad/Alcalde) qué secciones paralizan puntualmente los requerimientos de la misión.', 'Cumplido'],
        ['CL 008', 'Persistencia de progreso', 'El entorno del Menú principal sabe reconocer e importar una puntuación (estrellas) forjada en rondas previas o cierres bruscos.', 'Cumplido']
    ]
)

add_heading_custom(doc, '2.6.1.3.4.\tResultado del Sprint', level=3)
add_paragraph_custom(doc, 'Se formaron los verdaderos "músculos del descubrimiento". Ahora, en lugar de conectar rutas a la suerte o sin guion, se incita al estudiante a una reflexión constante. Valora una red, "solicita la revisión al sistema mediante analítica algorítmica" y experimenta el "fracaso pedagógico" identificando subgrafos defectuosos si falló. Solo logrará ser inscripto su logro permanente cuando alcance un arreglo rigurosamente conexo y apto. Ver')

# ========================
# SPRINT 4
# ========================
add_heading_custom(doc, '2.6.1.4.\tSprint 4', level=2)

add_heading_custom(doc, '2.6.1.4.1.\tObjetivo del sprint', level=3)
add_paragraph_custom(doc, 'Dotar a la vivencia de la calidad distintiva o cualidad "Juicy", aplicando efectos audiovisuales enrutativos inmersivos, destellos eléctricos lógicos de los cables y diseñando un epílogo triunfal GUI y sonoro que otorgue satisfacción tras cada sesión pesada y cognitiva al estudiante.')

add_heading_custom(doc, '2.6.1.4.2.\tPlanificación del sprint', level=3)
add_paragraph_custom(doc, 'Hablamos de la fase terminal del ciclo iterativo o Pulido Global Integrado. Se redactaron e instalaron Controladores Musicales interdependientes (ControladorAudio.client) de volumen dinámico, y el subsistema de Señalética Flotante (BillboardNombres guiadas por Tween) para dotar de modernismo los textos espaciales mientras se generaba el cortometraje minimalista de Cierre ("Flash GUI blanco de cálculo y tabla de puntaje"). La tabla 2.19 presenta las historias de usuario realizadas en el Sprint 4.')

add_paragraph_custom(doc, 'Tabla 2.19 Sprint Backlog - Sprint 4', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Descripción', 'Estimación (Horas)'],
    [
        ['CL 005', 'Audio adaptativo', 'ControladorAudio Cliente: Integrar melodías adaptativas con mecanismo Crossfading silenciado entre transiciones de ambiente, evitando estática repetitiva y cruda.', '10'],
        ['TG 08', 'Etiquetas de peso', 'Construir un motor de render de etiquetas o letreros paramétricos virtuales (Tag) que proyecten los diferentes "Pesos" (Costes de las líneas) visiblemente por sobre encima de la geometría opaca.', '12'],
        ['TG 09', 'Animación de cables', 'Simulador Direccional y de Velocidad (Beam/Tubos UV): Animación ininterrumpida direccional sobre las aristas, brindándole una visualidad fidedigna de sentido (a dónde fluye el poder eléctrico logístico) de un grafo dirigido.', '14'],
        ['CL 010', 'Pantalla de victoria', 'Cortinilla opaca de fin de partida que inmovilice el Input base deteniendo el reloj general del Gameplay procediendo a mostrar y tabular los "High Score" para coronación al usuario.', '8'],
        ['Total', '', '', '44']
    ]
)

add_heading_custom(doc, '2.6.1.4.3.\tRevisión del sprint', level=3)
add_paragraph_custom(doc, 'La Tabla 2.20 Revisión de criterios de aceptación - Sprint 4 presenta la retrospectiva del sprint 4, donde se presentan las historias de usuario y los criterios de aceptación verificados.')

add_paragraph_custom(doc, 'Tabla 2.20 Revisión de criterios de aceptación - Sprint 4', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Criterios de Aceptación', 'Estado'],
    [
        ['CL 005', 'Audio adaptativo', 'La métrica e inmersión sensorial no se detiene dolorosamente e impide saturaciones por picos abruptos del volumen al volver del Menú Principal.', 'Cumplido'],
        ['TG 08', 'Etiquetas de peso', 'Examinar y pasear en el 3D presenta fiel e instintivamente al avistador el coste respectivo (Ponderación/Weight) omitiendo obstáculos traslúcidos físicos.', 'Cumplido'],
        ['TG 09', 'Animación de cables', 'Las aristas logran lucir verdaderamente como filamentos brillantes asumiendo una velocidad UV observable desde el ánodo hacia el destino.', 'Cumplido'],
        ['CL 010', 'Pantalla de victoria', 'Pantalla visual de cese o GameOver / Triunfo infalible e inequívoco coronado con la acumulación temporal de todos los rubros económicos generados.', 'Cumplido']
    ]
)

add_heading_custom(doc, '2.6.1.4.4.\tResultado del Sprint', level=3)
add_paragraph_custom(doc, 'El proyecto superó el prototipado rudimentario y emergió siendo homólogo a una pieza sólida de educación entretenida del mercado comercial (Serious Game). Como resultado, la barrera inicial del estudiante, donde solían existir ansiedades y desestímulos provocados por la dureza abstracta del contenido de Estructura de Datos (EDA), ahora se contrarrestaba contundentemente con estímulos audiovisuales atractivos, dinámicos e inmersivos, culminando por completo la estructura técnica y lúdica del juego educativo planificado. Ver')

doc.save('Documentos/Sprints_2_3_4.docx')
print('Documento guardado en Documentos/Sprints_2_3_4.docx')
