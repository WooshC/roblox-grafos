from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return p

def add_paragraph_custom(doc, text, bold=False, italic=False, first_line_indent=Cm(1.25)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table_from_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = col_widths[i]
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            if col_widths and i < len(col_widths):
                row_cells[i].width = col_widths[i]
    
    doc.add_paragraph()
    return table

doc = Document()
section = doc.sections[0]
section.page_height = Cm(27.94)
section.page_width = Cm(21.59)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(2.54)

# ========================
# 2.5. Release Planning
# ========================
add_heading_custom(doc, '2.5.\tRelease Planning', level=2)
add_paragraph_custom(doc, 'La Tabla 2.12 Release Planning muestra cómo estas Historias (CL) y Mecánicas (TG) se agrupan lógicamente en Sprints para construir el juego paso a paso.')

add_paragraph_custom(doc, 'Tabla 2.12 Release Planning', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['Sprint 1: Fundamentos', 'Sprint 2: Lógica y Restricciones', 'Sprint 3: Validación, UI y Sistema'],
    [
        ['CL 001: Representación de nodos', 'TG 03: Eliminar conexiones', 'TG 08: Etiquetas de peso'],
        ['TG 01: Crear conexiones', 'TG 04: Presupuesto', 'TG 09: Matriz de adyacencia'],
        ['TG 02: Movimiento WASD', 'TG 05: Validación de reglas', 'TG 10: Ejecutar red'],
        ['CL 004: Construcción de grafos', 'TG 06: Tiempo límite', 'CL 002: Modo análisis'],
        ['', 'TG 07: Activación de nodos', 'CL 006: Simulación'],
        ['', '', 'CL 007: Visualización de costos'],
        ['', '', 'CL 008: Guardado y despliegue'],
        ['', '', 'CL 009: Gestión de estados (menú/reinicio)'],
    ]
)

# ========================
# 2.6.1.1. Sprint 1
# ========================
add_heading_custom(doc, '2.6.1.1.\tSprint 1', level=2)

# 2.6.1.1.1. Objetivo del sprint
add_heading_custom(doc, '2.6.1.1.1.\tObjetivo del sprint', level=3)
add_paragraph_custom(doc, 'Desarrollar el entorno tridimensional de "Villa Conexa", implementar la movilidad del personaje y asegurar la correcta representación visual de los nodos (postes) para sentar las bases de la interacción.')

# 2.6.1.1.2. Planificación del sprint
add_heading_custom(doc, '2.6.1.1.2.\tPlanificación del sprint', level=3)
add_paragraph_custom(doc, 'Para lograr el objetivo del Sprint 1 se implementaron las funciones de movimiento con teclado y ratón, se diseñó el escenario base con los activos gráficos de los nodos, y se desarrolló el "Controlador de Boot" que gestiona la máquina de estados entre el menú principal y el gameplay activo. La tabla 2.13 presenta las historias de usuario realizadas en el Sprint 1.')

add_paragraph_custom(doc, 'Tabla 2.13 Sprint Backlog - Sprint 1', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Descripción', 'Estimación (Horas)'],
    [
        ['CL 001', 'Representación de grafos', 'Visualización de nodos en el entorno 3D', '16'],
        ['TG 02', 'Movimiento WASD', 'Implementación de movimiento libre del jugador', '8'],
        ['TG 01', 'Crear conexiones', 'Sistema de clic para conectar nodos', '12'],
        ['CL 004', 'Construcción de grafos', 'Interacción general con nodos y aristas', '12'],
        ['Total', '', '', '48']
    ]
)

# 2.6.1.1.3. Revisión del sprint
add_heading_custom(doc, '2.6.1.1.3.\tRevisión del sprint', level=3)
add_paragraph_custom(doc, 'La Tabla 2.14 Revisión de criterios de aceptación - Sprint 1 presenta la retrospectiva del sprint 1, donde se presentan las historias de usuario y los criterios de aceptación verificados.')

add_paragraph_custom(doc, 'Tabla 2.14 Revisión de criterios de aceptación - Sprint 1', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Criterios de Aceptación', 'Estado'],
    [
        ['CL 001', 'Representación de grafos', 'Los nodos se visualizan correctamente en el entorno 3D y son identificables por el jugador', 'Cumplido'],
        ['TG 02', 'Movimiento WASD', 'El jugador puede desplazarse libremente por el escenario utilizando las teclas WASD sin errores de control', 'Cumplido'],
        ['TG 01', 'Crear conexiones', 'Al hacer clic en un nodo y luego en otro, se genera una conexión visual entre ambos', 'Cumplido'],
        ['CL 004', 'Construcción de grafos', 'El jugador puede interactuar con múltiples nodos y construir una red básica sin fallos', 'Cumplido']
    ]
)

# 2.6.1.1.4. Resultado del Sprint
add_heading_custom(doc, '2.6.1.1.4.\tResultado del Sprint', level=3)
add_paragraph_custom(doc, 'A continuación, se puede visualizar el resultado del Sprint 1, donde se observa al personaje principal en el entorno de Villa Conexa con los nodos desplegados listos para la interacción. Ver')

doc.save('Documentos/Capitulo_2_Sprints.docx')
print('Documento guardado en Documentos/Capitulo_2_Sprints.docx')
