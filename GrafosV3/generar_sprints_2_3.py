from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return p

def add_paragraph_custom(doc, text, bold=False, italic=False, first_line_indent=Cm(1.25)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table_from_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], 'D9E2F3')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = col_widths[i]
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            if col_widths and i < len(col_widths):
                row_cells[i].width = col_widths[i]
    
    doc.add_paragraph()
    return table

doc = Document()
section = doc.sections[0]
section.page_height = Cm(27.94)
section.page_width = Cm(21.59)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(2.54)

# ========================
# SPRINT 2
# ========================
add_heading_custom(doc, '2.6.1.2.\tSprint 2', level=2)

add_heading_custom(doc, '2.6.1.2.1.\tObjetivo del sprint', level=3)
add_paragraph_custom(doc, 'Implementar las reglas de interacción del grafo, permitiendo al jugador eliminar conexiones erróneas, gestionar un presupuesto limitado, validar las reglas topológicas del grafo, establecer un tiempo límite por nivel y activar o desactivar nodos según condiciones específicas de la misión.')

add_heading_custom(doc, '2.6.1.2.2.\tPlanificación del sprint', level=3)
add_paragraph_custom(doc, 'Para lograr el objetivo del Sprint 2 se desarrollaron los sistemas de validación de conexiones, el manejo de presupuesto, la eliminación de cables, el temporizador de niveles y la lógica de activación de nodos. Se implementaron los módulos encargados de verificar la legalidad de cada arista creada, controlar el gasto del presupuesto inicial, permitir el desmontaje de tramos incorrectos, limitar el tiempo de resolución por nivel y gestionar el estado activo o inactivo de cada nodo de la red. La tabla 2.15 presenta las historias de usuario y mecánicas de juego realizadas en el Sprint 2.')

add_paragraph_custom(doc, 'Tabla 2.15 Sprint Backlog - Sprint 2', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Descripción', 'Estimación (Horas)'],
    [
        ['TG 03', 'Eliminar conexiones', 'Crear una función de desmontaje controlado que barra tramos incorrectos suprimiéndolos del registro transaccional lógico de memoria.', '6'],
        ['TG 04', 'Sistema de presupuesto', 'Manejo contable del peso de las aristas, afectando de forma restada la variable económica base o presupuesto.', '10'],
        ['TG 05', 'Validación de reglas', 'Asegurar la validación topológica negando flujos ilegales o superposición reiterativa de la misma arista.', '10'],
        ['TG 06', 'Tiempo límite', 'Establecer un temporizador por nivel que limite el tiempo disponible para completar la conexión de la red.', '8'],
        ['TG 07', 'Activación de nodos', 'Implementar la lógica de activación y desactivación de nodos según condiciones específicas de cada misión.', '10'],
        ['Total', '', '', '44']
    ]
)

add_heading_custom(doc, '2.6.1.2.3.\tRevisión del sprint', level=3)
add_paragraph_custom(doc, 'La Tabla 2.16 Revisión de criterios de aceptación - Sprint 2 presenta la retrospectiva del sprint 2, donde se presentan las historias de usuario y los criterios de aceptación verificados.')

add_paragraph_custom(doc, 'Tabla 2.16 Revisión de criterios de aceptación - Sprint 2', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Criterios de Aceptación', 'Estado'],
    [
        ['TG 03', 'Eliminar conexiones', 'El jugador puede eliminar una conexión existente y esta desaparece tanto visual como lógicamente de la matriz de adyacencia.', 'Cumplido'],
        ['TG 04', 'Sistema de presupuesto', 'El presupuesto inicial se decrementa correctamente al crear conexiones y el juego impide acciones que generen saldo negativo.', 'Cumplido'],
        ['TG 05', 'Validación de reglas', 'El sistema valida que no se creen aristas duplicadas ni conexiones que violen las reglas de direccionalidad definidas por el nivel.', 'Cumplido'],
        ['TG 06', 'Tiempo límite', 'El temporizador se inicia al comenzar el nivel, se visualiza en el HUD y finaliza la partida al llegar a cero.', 'Cumplido'],
        ['TG 07', 'Activación de nodos', 'Los nodos responden correctamente a los eventos de activación y desactivación según las condiciones programadas en cada misión.', 'Cumplido']
    ]
)

add_heading_custom(doc, '2.6.1.2.4.\tResultado del Sprint', level=3)
add_paragraph_custom(doc, 'Se consolidó el núcleo de la jugabilidad estratégica. El estudiante ya no solo observa el grafo, sino que interactúa con él bajo restricciones reales: presupuesto finito, tiempo limitado y reglas topológicas estrictas. Cada acción de conexión o eliminación actualiza la matriz de adyacencia en tiempo real, proporcionando retroalimentación inmediata sobre la legalidad y el costo de la operación. Ver')

# ========================
# SPRINT 3
# ========================
add_heading_custom(doc, '2.6.1.3.\tSprint 3', level=2)

add_heading_custom(doc, '2.6.1.3.1.\tObjetivo del sprint', level=3)
add_paragraph_custom(doc, 'Dotar al juego de herramientas de visualización avanzada de la teoría de grafos, modos de análisis algorítmico, simulación del flujo de red, persistencia de progreso del jugador y gestión completa del ciclo de vida del juego, integrando menú principal, reinicio de niveles y transiciones de estado.')

add_heading_custom(doc, '2.6.1.3.2.\tPlanificación del sprint', level=3)
add_paragraph_custom(doc, 'Para lograr el objetivo del Sprint 3 se implementaron la visualización de pesos en las conexiones, la matriz de adyacencia interactiva, el ejecutor y validador de redes, el modo de análisis con algoritmos BFS/DFS, la simulación del flujo eléctrico, el despliegue web, el guardado de progreso mediante Roblox DataStore y la gestión integral de estados del juego. Se desarrollaron los paneles de inspección visual y las herramientas diagnósticas que permiten al estudiante verificar la conectividad de su red antes de finalizar el nivel. La tabla 2.17 presenta las historias de usuario y mecánicas de juego realizadas en el Sprint 3.')

add_paragraph_custom(doc, 'Tabla 2.17 Sprint Backlog - Sprint 3', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Descripción', 'Estimación (Horas)'],
    [
        ['TG 08', 'Visualización de pesos', 'Mostrar numéricamente el peso o costo de cada arista sobre la conexión visual en el entorno 3D.', '6'],
        ['TG 09', 'Matriz de adyacencia', 'Construir y mantener actualizada en tiempo real la matriz de adyacencia que representa el estado actual del grafo.', '12'],
        ['TG 10', 'Ejecutar/validar red', 'Implementar el motor de validación que ejecuta algoritmos de recorrido (BFS/DFS) para verificar conectividad y cumplimiento de misiones.', '14'],
        ['CL 002', 'Modo análisis y visualización de algoritmos', 'Habilitar una vista especializada que ilustre paso a paso la ejecución de los algoritmos de grafos sobre la red construida.', '14'],
        ['CL 006', 'Ejecución y simulación de la red', 'Simular el flujo de energía o datos a través de las conexiones establecidas, mostrando visualmente el recorrido completo.', '14'],
        ['CL 007', 'Visualización de costos', 'Integrar en el HUD la visualización detallada de los costos acumulados, restantes y por arista.', '8'],
        ['CL 008', 'Guardado y despliegue', 'Persistir el progreso del jugador (niveles desbloqueados, puntuaciones, tiempos) mediante Roblox DataStore.', '14'],
        ['CL 009', 'Gestión del ciclo de vida', 'Gestionar las transiciones entre menú principal, selección de nivel, gameplay activo, pausa y reinicio sin perdida de estado.', '18'],
        ['Total', '', '', '100']
    ]
)

add_heading_custom(doc, '2.6.1.3.3.\tRevisión del sprint', level=3)
add_paragraph_custom(doc, 'La Tabla 2.18 Revisión de criterios de aceptación - Sprint 3 presenta la retrospectiva del sprint 3, donde se presentan las historias de usuario y los criterios de aceptación verificados.')

add_paragraph_custom(doc, 'Tabla 2.18 Revisión de criterios de aceptación - Sprint 3', bold=True, first_line_indent=Cm(0))
add_table_from_data(doc,
    ['ID', 'Historia / Mecánica', 'Criterios de Aceptación', 'Estado'],
    [
        ['TG 08', 'Visualización de pesos', 'Cada conexión visible en el escenario muestra su peso o costo asociado de forma clara y legible para el jugador.', 'Cumplido'],
        ['TG 09', 'Matriz de adyacencia', 'La matriz refleja fielmente el estado actual del grafo, actualizándose automáticamente al crear o eliminar aristas.', 'Cumplido'],
        ['TG 10', 'Ejecutar/validar red', 'El sistema detecta correctamente nodos aislados, ciclos prohibidos y verifica el cumplimiento de los objetivos de conectividad.', 'Cumplido'],
        ['CL 002', 'Modo análisis y visualización de algoritmos', 'El modo análisis muestra paso a paso la ejecución de BFS/DFS sobre la red, resaltando nodos visitados y aristas recorridas.', 'Cumplido'],
        ['CL 006', 'Ejecución y simulación de la red', 'La simulación de flujo recorre todas las aristas conectadas desde el nodo origen hasta los destinos sin interrupciones visuales.', 'Cumplido'],
        ['CL 007', 'Visualización de costos', 'El HUD actualiza en tiempo real el presupuesto disponible, el gasto acumulado y el costo individual de cada arista seleccionada.', 'Cumplido'],
        ['CL 008', 'Guardado y despliegue', 'El progreso del jugador persiste entre sesiones, incluyendo niveles completados, estrellas obtenidas y mejores tiempos.', 'Cumplido'],
        ['CL 009', 'Gestión del ciclo de vida', 'Las transiciones entre menú, nivel y gameplay son fluidas, sin pérdida de datos ni estados inconsistentes al reiniciar o cambiar de nivel.', 'Cumplido']
    ]
)

add_heading_custom(doc, '2.6.1.3.4.\tResultado del Sprint', level=3)
add_paragraph_custom(doc, 'El juego alcanzó su madurez funcional. El estudiante dispone ahora de un entorno completo donde puede construir, analizar y simular grafos en tres dimensiones, con retroalimentación algorítmica inmediata y persistencia de su trayectoria de aprendizaje. La integración del menú principal, el guardado de progreso y las herramientas de análisis convierten al prototipo en una experiencia educativa cerrada y reusable. Ver')

doc.save('Documentos/Sprints_2_3.docx')
print('Documento guardado en Documentos/Sprints_2_3.docx')
